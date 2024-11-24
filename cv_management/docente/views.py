from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Docente
from django.http import HttpResponse #wilson
from docx import Document #wilson
from docx.oxml import OxmlElement#wilson
from docx.oxml.ns import qn #wilson
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from django.contrib.auth.decorators import login_required
from django.contrib.auth.hashers import check_password
from django.contrib.auth import authenticate, login

@login_required
def update_teacher(request):
    if request.method == 'POST':
        # Obtén los datos del formulario
        nombre = request.POST.get('nombre')
        apellido = request.POST.get('apellido')
        cedula = request.POST.get('cedula')
        num_telefono = request.POST.get('num_telefono')
        universidad = request.POST.get('universidad')
        cargo = request.POST.get('cargo')
        correo = request.POST.get('correo')
        facultad = request.POST.get('facultad')
        categoria = request.POST.get('categoria')
        especialidad = request.POST.get('especialidad')
        tipo_contrato = request.POST.get('tipo_contrato')
        fecha_contratacion = request.POST.get('fecha_contratacion')

        user = request.user
        # Verifica si el correo ya está registrado en Docente
        try:
            docente = Docente.objects.get(correo=user.email)
        except Docente.DoesNotExist:
            docente = None
        
        if docente:
            # Si el docente ya existe, solo actualizamos sus datos
            docente.nombre = nombre
            docente.apellido = apellido
            docente.cedula = cedula
            docente.num_telefono = num_telefono
            docente.universidad = universidad
            docente.cargo = cargo
            docente.correo = correo
            docente.facultad = facultad
            docente.categoria = categoria
            docente.especialidad = especialidad
            docente.tipo_contrato = tipo_contrato
            docente.fecha_contratacion = fecha_contratacion
            docente.estado = 0
            docente.save()

            # Actualizamos también los datos del usuario relacionado
            user.first_name = nombre
            user.last_name = apellido
            user.email = correo  # Actualiza el correo si es necesario
            user.save()

            messages.success(request, "Datos del docente actualizados con éxito.")
        return redirect('update')

    return render(request, 'update.html')  # La plantilla del formulario de actualización

@login_required
def update_picture(request):
    if request.method == "POST":
        docente = request.user.docente
        print(f"Docente: {docente}")
        # Verificamos si se ha subido una imagen
        photo = request.FILES.get('profile_picture')
        if photo:
            print(f"Uploading photo: {photo}")
            docente.foto = photo
            docente.save()
            return redirect('update_teacher')  # Redirige a una página de éxito o de actualización
        
        # Lógica para eliminar la imagen de perfil si se envía la solicitud para eliminar
        if request.POST.get('remove_picture') == 'true':
            docente.foto.delete(save=True)
            return redirect('update_teacher')  # Redirige después de eliminar la imagen

    # Si no es un POST o no se subió ningún archivo, puedes redirigir o mostrar un mensaje
    return HttpResponse("Método no permitido o no se ha subido una imagen", status=400)

@login_required
def settings_teacher(request):
    return render(request, 'settings_teacher.html')

@login_required
def reset_password_teacher(request):
    if request.method == 'POST':
        current_password = request.POST.get('current-password')
        new_password = request.POST.get('new-password')
        
        user = request.user
        
        # Verifica si la contraseña actual es correcta utilizando check_password
        if check_password(current_password, user.password):
            # Si la contraseña actual es correcta, actualiza la nueva contraseña
            user.set_password(new_password)
            user.save()
            messages.success(request, "Contraseña actualizada con éxito.")
            login(request, user)   
        else:
            messages.error(request, "La contraseña actual es incorrecta.")
        
        return redirect('settings_teacher')
    
    return render(request, 'settings_teacher.html')  # La plantilla del formulario de registro

#wilson
@login_required
def generate_curriculum(request, docente_id):
    # Crear el documento en formato Word para edición
    docente = get_object_or_404(Docente, id=docente_id)
    doc = Document()

    # Títulos y subtítulos
    doc.add_heading(f'Currículum de {docente.nombre} {docente.apellido}', level=1)

    # Información de contacto
    doc.add_heading("Información de Contacto", level=2)
    doc.add_paragraph(f"Cédula: {docente.cedula or 'N/A'}")
    doc.add_paragraph(f"Correo: {docente.correo or 'N/A'}")
    doc.add_paragraph(f"Teléfono: {docente.num_telefono or 'N/A'}")

    # Información profesional
    doc.add_heading("Información Profesional", level=2)
    doc.add_paragraph(f"Universidad: {docente.universidad or 'N/A'}")
    doc.add_paragraph(f"Facultad: {docente.facultad or 'N/A'}")
    doc.add_paragraph(f"Especialidad: {docente.especialidad or 'N/A'}")
    doc.add_paragraph(f"Categoría: {docente.categoria or 'N/A'}")

    # Detalles del contrato
    doc.add_heading("Detalles del Contrato", level=2)
    doc.add_paragraph(f"Tipo de Contrato: {docente.tipo_contrato or 'N/A'}")
    doc.add_paragraph(f"Estado: {docente.estado or 'N/A'}")
    doc.add_paragraph(f"Fecha de Contratación: {docente.fecha_contratacion.strftime('%d/%m/%Y') if docente.fecha_contratacion else 'N/A'}")

     # Guardar documento Word temporal
    # temp_word = BytesIO()
    # doc.save(temp_word)
    # temp_word.seek(0)

    # # Convertir el documento en PDF
    # response = HttpResponse(content_type='application/pdf')
    # response['Content-Disposition'] = f'attachment; filename="{docente.nombre}_{docente.apellido}_curriculum.pdf"'

    # # Crear el PDF con ReportLab
    # pdf_canvas = canvas.Canvas(response, pagesize=letter)
    # pdf_canvas.setFont("Helvetica", 12)

    # Convertir el documento en PDF y almacenarlo en memoria
    pdf_buffer = BytesIO()
    pdf_canvas = canvas.Canvas(pdf_buffer, pagesize=letter)
    pdf_canvas.setFont("Helvetica", 12)

    # Añadir contenido al PDF
    y_position = 750  # Posición vertical inicial
    pdf_canvas.drawString(100, y_position, f"Currículum de {docente.nombre} {docente.apellido}")
    y_position -= 30
    pdf_canvas.drawString(100, y_position, "Información de Contacto:")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Cédula: {docente.cedula or 'N/A'}")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Correo: {docente.correo or 'N/A'}")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Teléfono: {docente.num_telefono or 'N/A'}")
    y_position -= 40

    pdf_canvas.drawString(100, y_position, "Información Profesional:")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Universidad: {docente.universidad or 'N/A'}")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Facultad: {docente.facultad or 'N/A'}")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Especialidad: {docente.especialidad or 'N/A'}")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Categoría: {docente.categoria or 'N/A'}")
    y_position -= 40

    pdf_canvas.drawString(100, y_position, "Detalles del Contrato:")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Tipo de Contrato: {docente.tipo_contrato or 'N/A'}")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Estado: {docente.estado or 'N/A'}")
    y_position -= 20
    pdf_canvas.drawString(120, y_position, f"Fecha de Contratación: {docente.fecha_contratacion.strftime('%d/%m/%Y') if docente.fecha_contratacion else 'N/A'}")

    # Finalizar PDF
    pdf_canvas.save()
    pdf_buffer.seek(0)

    response = HttpResponse(pdf_buffer, content_type='application/pdf')
    return response

